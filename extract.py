import docx

def extract():
    doc = docx.Document("RLT REPORT.docx")
    with open("extracted.txt", "w", encoding="utf-8") as f:
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                f.write(f"{i}: {p.text}\n")
        f.write("\n--- TABLES ---\n")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        f.write(f"Table {t_idx} - {r_idx},{c_idx}: {cell.text.strip()}\n")

if __name__ == "__main__":
    extract()
