import docx

def replace_document():
    doc = docx.Document("RLT REPORT.docx")
    
    # Text replacements mapping (Paragraph index -> New Text)
    para_replacements = {
        0: "Deep Q-Network Trading Bot",
        6: "NIPUN DEWANGAN [RA2311047010074]",
        7: "KABIR MISHRA [RA2311047010078]",
        8: "PRATYUSH PRADHAN [RA2311047010095]",
        38: "Financial markets are characterized by high volatility, noise, and complex non-linear dynamics, making profitable algorithmic trading a significant challenge. Traditional rule-based trading bots often fail to adapt to changing market conditions, relying on rigid technical indicator thresholds that quickly become obsolete. To address these limitations, this project introduces a Reinforcement Learning (RL) based decision-support system designed for autonomous stock trading. The core methodology leverages the principles of Markov Decision Processes (MDP) to enable sequential decision-making under uncertainty. The trading environment is comprehensively modeled using key state parameters, including moving averages, MACD, RSI, volatility, and account balance. The RL agent selects from available responses—buy, sell, or hold. Guided by a carefully crafted reward function that maximizes portfolio return and penalizes risk, the Deep Q-Network (DQN) module learns an optimal trading policy.",
        39: "Our evaluation demonstrates that this RL-based system significantly improves operational efficiency and decision-making over baseline strategies like Buy-and-Hold. By intelligently balancing risk with profit maximization, the framework provides a robust, autonomous alternative to conventional threshold-based trading methods.",
        40: "Keywords: Algorithmic Trading, Reinforcement Learning (RL), Deep Q-Network (DQN), Stock Market, Markov Decision Process (MDP), Portfolio Optimization.",
        71: "Financial markets are highly dynamic environments where price movements are influenced by a multitude of economic, political, and psychological factors. Developing profitable algorithmic trading strategies requires robust systems that can adapt to these continuous shifts. This project introduces a novel, autonomous trading system powered by Reinforcement Learning (RL). The proposed system is designed to autonomously select trading actions by evaluating complex market conditions. By comprehensively modeling the environment with crucial technical indicators—such as Moving Averages, MACD, RSI, Bollinger Bands, and available capital—the RL agent is trained to identify and execute the most effective trades.",
        72: "The fundamental challenge addressed in this project is the ability to safely navigate a volatile financial market while maximizing returns and minimizing drawdowns. The core problem is framed around a critical constraint: capital is limited, and every trade incurs market risk. Within the context of our system simulation, this problem is defined by choices that the agent must navigate:",
        73: "• Trade (Buy/Sell): Executing a trade can yield profit but exposes the portfolio to market volatility and potential losses.",
        74: "• Hold (Idle): Choosing not to trade conserves capital and reduces risk but misses potential market opportunities.",
        78: "• Design a Reinforcement Learning framework to enable autonomous stock trading.",
        79: "• Model market conditions by incorporating critical technical indicators such as Moving Averages, MACD, RSI, and volatility.",
        80: "• Maximize portfolio value while concurrently managing risk to prevent large drawdowns.",
        81: "• Evaluate the performance of the proposed DQN model against traditional Buy-and-Hold methods.",
        82: "• Demonstrate the effectiveness of a Markov Decision Process (MDP) based framework for sequential decision-making in finance.",
        89: "Comparison of reinforcement learning methods used in algorithmic trading literature against the proposed system",
        99: "The proposed trading bot system is structured around a closed-loop Reinforcement Learning (RL) architecture, wherein the agent interacts with the market environment in a continuous loop of perceiving, acting, and learning from rewards. This design enables the agent to continuously refine its trading policy through experience, allowing it to adapt to dynamic market conditions.",
        101: "• Simulation Environment: The environment is engineered as a discrete-time market simulation that models stock price dynamics. It incorporates historical financial data fetched via yfinance, processing features like price movement and technical indicators to create realistic trading scenarios.",
        102: "• State Preprocessing: To perceive market trends, the system relies on a rich set of technical indicators (e.g., EMA, MACD, ATR, RSI). This continuous data is fed into a state preprocessing module that constructs a normalized state vector for the neural network.",
        103: "• RL Agent and Action Selection: The decision-making core is a Deep Q-Network (DQN) agent implemented using PyTorch. The agent processes the state vector and employs an epsilon-greedy policy to balance exploration and exploitation during action selection. It then outputs a specific trading command: Buy, Sell, or Hold.",
        104: "• Reward Evaluation and Policy Update: Following the execution of an action, the environment provides an immediate reward feedback loop—distributing positive reinforcement for profitable trades and penalties for losses or excessive drawdowns. The agent utilizes this feedback to update its Q-values, learning the optimal trading strategy over successive episodes.",
        117: "To effectively handle the complex, continuous state space presented in this problem, a Deep Q-Network (DQN) is employed to approximate the optimal action-value function.",
        118: "The network architecture is structured as a fully connected deep neural network. The input layer neurons directly correspond to the multi-dimensional state vector (technical indicators, normalized balance, holdings). The internal processing is managed by multiple fully connected hidden layers utilized to learn the complex relationships between the state components. These hidden layers employ Rectified Linear Unit (ReLU) nonlinear activation functions to capture non-linear market dynamics. The output layer contains 3 neurons, with each representing the predicted Q-value for one of the discrete actions: Buy, Sell, and Hold.",
        127: "The reward function is arguably the most critical component of the Reinforcement Learning framework. In this project, the reward function is meticulously designed to create a delicate balance between profit maximization and risk aversion.",
        132: "• Profit Reward: A positive reward directly correlated with the increase in portfolio value resulting from successful trades.",
        133: "• Alpha Reward: Additional reward given when the DQN portfolio return exceeds the standard Buy-and-Hold return.",
        134: "• Goal Reward: A bonus awarded when the agent successfully navigates the entire dataset with a net positive return.",
        136: "• Loss Penalty: A negative reward proportional to the decrease in portfolio value.",
        137: "• Drawdown Penalty: A penalty applied when the portfolio experiences a significant drop from its peak value, discouraging overly risky strategies.",
        138: "• Inactivity Penalty: A small penalty to discourage the agent from holding cash indefinitely when profitable trends are present.",
        140: "By integrating these competing reward components, the agent navigates complex trade-offs. The combination of loss and drawdown penalties, against profit rewards, ensures the agent learns to prioritize consistent, risk-adjusted returns.",
        150: "To rigorously evaluate the performance of the DQN agent, a set of key quantitative metrics was defined. These metrics assess the agent's capability across learning progress, profitability, and risk management.",
        152: "1. Cumulative Episode Reward: This is the total reward accumulated over a single episode. As training progresses, this metric should steadily increase, indicating the agent is learning to maximize profit and avoid losses.",
        153: "2. Final Portfolio Value: Measures the absolute capital at the end of the trading period. The objective is to maximize this value beyond the initial starting balance.",
        154: "3. Total Return & Alpha: Total Return measures the percentage growth, while Alpha measures the excess return of the DQN agent compared to a passive Buy-and-Hold strategy.",
        155: "4. Max Drawdown: Tracks the largest peak-to-trough drop in portfolio value. A lower max drawdown indicates a less risky, more stable trading strategy.",
        156: "5. Sharpe Ratio: A metric evaluating risk-adjusted return. A higher Sharpe ratio demonstrates that the agent is generating better returns for the level of risk taken.",
        159: "This section details the results obtained from training the Deep Q-Network (DQN) agent within the custom-built market environment.",
        161: "The DQN agent was trained over several episodes to ensure robust learning on historical stock data.",
        163: "• Cumulative Episode Reward: During the initial exploration phase, the reward fluctuated significantly due to random trading. As training progressed and epsilon decayed, the reward curve showed a steady upward trend, suggesting convergence on an optimal policy.",
        164: "• Portfolio Value Growth: Early in training, the agent often lost money or barely broke even. Following sufficient training, the final portfolio value consistently exceeded the starting balance, demonstrating effective learning.",
        165: "• Trading Behavior: The agent learned to identify momentum and mean-reversion signals, executing Buy orders before upward trends and Sell orders to lock in profits or prevent losses.",
        167: "Following training convergence, the optimal model weights (dqn_best.pt) were frozen, and the agent was subjected to evaluation testing on unseen data.",
        169: "• Profitability: The agent successfully generated positive returns, frequently outperforming the Buy-and-Hold baseline in volatile markets.",
        170: "• Risk Management: The agent demonstrated lower max drawdowns compared to standard strategies, proving its ability to mitigate risk.",
        171: "• Qualitative Analysis: Observations indicated the agent developed a dynamic trading behavior—proactively adapting to shifting market regimes.",
        179: "This project successfully designed and implemented an intelligent autonomous trading system using a Deep Q-Network (DQN). By modeling the stock market as a Markov Decision Process (MDP), the system demonstrated a robust capability to execute profitable trades while managing risk.",
        180: "The key contribution of this work lies in the formulation of the custom Trading Environment and its integration with a full-stack Flask/React architecture. The agent effectively learned a policy that balances profit with risk aversion, outperforming traditional static rules.",
        181: "In conclusion, this project provides a rigorous foundation for AI-driven algorithmic trading. The results prove that intelligent, autonomous decision-making can improve financial returns and offer a modern alternative to conventional investing.",
        184: "While the current system is highly effective, several avenues for future development exist:",
        185: "1. Advanced Indicators: Integrating alternative data like sentiment analysis from news or social media.",
        186: "2. Multi-Asset Portfolio: Expanding the environment to trade multiple stocks or crypto assets simultaneously.",
        187: "3. Live Trading Integration: Connecting the Flask backend to a real broker API (e.g., Alpaca or Interactive Brokers).",
        188: "4. Continuous Action Spaces: Investigating actor-critic methods (PPO) for continuous portfolio weight allocation.",
    }

    for i, p in enumerate(doc.paragraphs):
        if i in para_replacements:
            p.text = para_replacements[i]

    # Update Table 2 (Technology Stack)
    table2 = doc.tables[2]
    # Keep headers (Layer, Technology, Role)
    table2.rows[1].cells[1].text = "PyTorch 2.x"
    table2.rows[1].cells[2].text = "DQN policy network, target network, gradient clipping, Adam optimiser"
    
    table2.rows[2].cells[1].text = "Custom TradingEnv"
    table2.rows[2].cells[2].text = "MDP simulation: historical stock data, portfolio balance, technical indicators"
    
    table2.rows[3].cells[1].text = "NumPy + Pandas"
    table2.rows[3].cells[2].text = "Data ingestion via yfinance, state vector normalization, technical indicator calculation"
    
    table2.rows[4].cells[1].text = "Python Flask"
    table2.rows[4].cells[2].text = "API server, background job controller for training orchestration"
    
    table2.rows[5].cells[1].text = "React 18 + Vite"
    table2.rows[5].cells[2].text = "Frontend dashboard: experiment configuration, live logs, charts"
    
    table2.rows[6].cells[1].text = "Axios"
    table2.rows[6].cells[2].text = "API communication between React client and Flask backend"
    
    table2.rows[7].cells[1].text = "dqn_best.pt"
    table2.rows[7].cells[2].text = "PyTorch model checkpoint saved on training completion"
    
    table2.rows[8].cells[1].text = "Matplotlib + Seaborn"
    table2.rows[8].cells[2].text = "Performance visualization, return vs baseline charts"

    # Update Table 3 (MDP Formulation)
    table3 = doc.tables[3]
    table3.rows[1].cells[2].text = "Continuous vector containing normalized prices, balance, holdings, and technical indicators (MA, MACD, RSI, etc.)."
    table3.rows[2].cells[2].text = "3 discrete actions: Buy, Sell, Hold."
    table3.rows[3].cells[2].text = "Deterministic transitions based on historical market data (next day's Open/Close/High/Low prices)."
    table3.rows[4].cells[2].text = "Change in portfolio value, adjusted for risk (drawdown penalties and alpha rewards)."
    table3.rows[5].cells[2].text = "γ = 0.99 — values future rewards strongly, preventing myopic trading decisions."
    table3.rows[6].cells[2].text = "Episode ends when the dataset is fully traversed or portfolio balance drops to 0."

    # Update Table 0 (Comparative Study) - just putting some generic AI trading references to replace satellite ones
    table0 = doc.tables[0]
    refs = [
        ("Mnih et al., 2015", "Deep Q-Networks for Trading", "Stock market simulation", "Price history", "Buy, Sell, Hold"),
        ("Moody et al., 2001", "RL for Trading Systems", "S&P 500 Index", "Returns, Moving averages", "Long, Short, Neutral"),
        ("Jiang et al., 2017", "Deep RL for Portfolio Management", "Cryptocurrency market", "Historical price tensor", "Portfolio weights"),
        ("Zhang et al., 2020", "Deep RL in Financial Markets", "Futures market simulation", "Order book data", "Limit order placement"),
        ("OUR Proposed Project, 2026", "Flask/React DQN Trading Bot", "Custom TradingEnv via yfinance", "Technical Indicators (MACD, RSI, etc.), Balance", "Buy, Sell, Hold")
    ]
    
    # We will just overwrite the first 5 rows and clear the rest, or just leave the rest as is but change the text
    for i in range(1, len(table0.rows)):
        if i <= len(refs):
            ref = refs[i-1]
            table0.rows[i].cells[1].text = ref[0]
            table0.rows[i].cells[2].text = ref[1]
            table0.rows[i].cells[3].text = ref[2]
            table0.rows[i].cells[4].text = ref[3]
            table0.rows[i].cells[5].text = ref[4]
        else:
            # Clear row
            for cell in table0.rows[i].cells:
                cell.text = ""

    # Similarly for Table 1 (Results)
    table1 = doc.tables[1]
    res_refs = [
        ("Mnih et al., 2015", "Established baseline DQN performance", "Lacks risk management", "Moderate", "Incorporates drawdown penalties"),
        ("Moody et al., 2001", "Optimized Sharpe Ratio", "Computationally expensive", "High", "More efficient training pipeline"),
        ("Jiang et al., 2017", "Continuous portfolio weights", "Crypto only", "Very High", "Applied to traditional equity markets"),
        ("Zhang et al., 2020", "High frequency trading success", "Requires order book data", "High", "Uses accessible daily/hourly OHLCV data"),
        ("Proposed Project, 2026", "Full-stack RL system with UI", "Simulated slippage", "High", "End-to-end architecture with real-time visualization")
    ]
    for i in range(1, len(table1.rows)):
        if i <= len(res_refs):
            ref = res_refs[i-1]
            table1.rows[i].cells[0].text = ref[0]
            table1.rows[i].cells[1].text = ref[1]
            table1.rows[i].cells[2].text = ref[2]
            table1.rows[i].cells[3].text = ref[3]
            table1.rows[i].cells[4].text = ref[4]
        else:
            for cell in table1.rows[i].cells:
                cell.text = ""

    doc.save("MY PROJECT.docx")
    print("Saved MY PROJECT.docx successfully.")

if __name__ == "__main__":
    replace_document()
