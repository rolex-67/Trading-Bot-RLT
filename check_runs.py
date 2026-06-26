import docx

def check_runs():
    doc = docx.Document("RLT REPORT.docx")
    
    print("Paragraph 38 runs:")
    for r in doc.paragraphs[38].runs:
        print(f" - '{r.text}' (bold: {r.bold})")
        
    print("\nParagraph 71 runs:")
    for r in doc.paragraphs[71].runs:
        print(f" - '{r.text}' (bold: {r.bold})")
        
    print("\nParagraph 72 runs:")
    for r in doc.paragraphs[72].runs:
        print(f" - '{r.text}' (bold: {r.bold})")

if __name__ == "__main__":
    check_runs()
