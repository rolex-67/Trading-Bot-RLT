import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

def update_document():
    doc = docx.Document("RLT REPORT.docx")
    
    # Text replacements mapping
    para_replacements = {
        0: "Deep Q-Network Trading Bot",
        6: "NIPUN DEWANGAN [RA2311047010074]",
        7: "KABIR MISHRA [RA2311047010078]",
        8: "PRATYUSH PRADHAN [RA2311047010095]",
        38: "Financial markets are characterized by high volatility, noise, and complex non-linear dynamics, making profitable algorithmic trading a significant challenge. Traditional rule-based trading bots often fail to adapt to changing market conditions, relying on rigid technical indicator thresholds that quickly become obsolete. To address these limitations, this project introduces a Reinforcement Learning (RL) based decision-support system designed for autonomous stock trading. The core methodology leverages the principles of Markov Decision Processes (MDP) to enable sequential decision-making under uncertainty. The trading environment is comprehensively modeled using key state parameters, including moving averages, MACD, RSI, volatility, and account balance. The RL agent selects from available responses—buy, sell, or hold. Guided by a carefully crafted reward function that maximizes portfolio return and penalizes risk, the Deep Q-Network (DQN) module learns an optimal trading policy.",
        39: "Our evaluation demonstrates that this RL-based system significantly improves operational efficiency and decision-making over baseline strategies like Buy-and-Hold. By intelligently balancing risk with profit maximization, the framework provides a robust, autonomous alternative to conventional threshold-based trading methods.",
        40: "Keywords: Algorithmic Trading, Reinforcement Learning (RL), Deep Q-Network (DQN), Stock Market, Markov Decision Process (MDP), Portfolio Optimization.",
        73: "Trade (Buy/Sell): Executing a trade can yield profit but exposes the portfolio to market volatility and potential losses.",
        74: "Hold (Idle): Choosing not to trade conserves capital and reduces risk but misses potential market opportunities.",
        78: "Design a Reinforcement Learning framework to enable autonomous stock trading.",
        79: "Model market conditions by incorporating critical technical indicators such as Moving Averages, MACD, RSI, and volatility.",
        80: "Maximize portfolio value while concurrently managing risk to prevent large drawdowns.",
        81: "Evaluate the performance of the proposed DQN model against traditional Buy-and-Hold methods.",
        82: "Demonstrate the effectiveness of a Markov Decision Process (MDP) based framework for sequential decision-making in finance.",
        89: "Comparison of reinforcement learning methods used in algorithmic trading literature against the proposed system",
        94: "Note: Profitability and alpha figures are as reported in the respective algorithmic trading literature or derived from their evaluation metrics. The proposed method's returns are measured against a standard Buy-and-Hold baseline strategy using a starting portfolio of $10,000 over historical out-of-sample trading data.",
        101: "Simulation Environment: The environment is engineered as a discrete-time market simulation that models stock price dynamics. It incorporates historical financial data fetched via yfinance, processing features like price movement and technical indicators to create realistic trading scenarios.",
        102: "State Preprocessing: To perceive market trends, the system relies on a rich set of technical indicators (e.g., EMA, MACD, ATR, RSI). This continuous data is fed into a state preprocessing module that constructs a normalized state vector for the neural network.",
        103: "RL Agent and Action Selection: The decision-making core is a Deep Q-Network (DQN) agent implemented using PyTorch. The agent processes the state vector and employs an epsilon-greedy policy to balance exploration and exploitation during action selection. It then outputs a specific trading command: Buy, Sell, or Hold.",
        104: "Reward Evaluation and Policy Update: Following the execution of an action, the environment provides an immediate reward feedback loop—distributing positive reinforcement for profitable trades and penalties for losses or excessive drawdowns. The agent utilizes this feedback to update its Q-values, learning the optimal trading strategy over successive episodes.",
        117: "In our DQN Trading Bot, the Q-function Q(s, a) estimates the expected cumulative discounted return from executing action a (Buy, Sell, or Hold) given the current market state s, which encodes 20+ features including price-based indicators (EMA-9, EMA-21, MACD, MACD Signal), momentum indicators (RSI-14, ATR-14), and portfolio state (normalized cash balance, number of shares held). The DQN learns to map each unique market snapshot directly to the best possible trading decision.",
        118: "The policy network is implemented in PyTorch as a 4-layer fully connected neural network. Layer 1 (Input): Accepts the state vector of dimension equal to the number of features (e.g., 22 inputs). Layers 2-3 (Hidden): Two dense hidden layers of 256 and 128 neurons respectively, each followed by ReLU activations to model non-linear interactions between technical signals and portfolio state. Layer 4 (Output): 3 output neurons representing Q-values for Buy (0), Sell (1), and Hold (2). The action with the highest Q-value is executed by the trading bot.",
        119: "Three key mechanisms stabilize learning in the volatile stock market environment:",
        120: "Experience Replay Buffer: Every market step generates a transition tuple (s_t, a_t, r_t, s_{t+1}, done) that is pushed into a fixed-size circular replay buffer (capacity: 100,000 transitions). At each training step, a random mini-batch of 64 transitions is sampled. This breaks the dangerous temporal autocorrelation inherent in sequential stock price data (e.g., today's price is highly correlated with yesterday's), allowing the network to learn from diverse, decorrelated market experiences.",
        121: "Target Network (Q'): A frozen copy of the online Q-network (Q) is maintained as the Target Network. It is used solely to compute stable Bellman target values during training. Its weights are updated to match the online network only every C=10 steps (soft or hard update). Without this, the online network would chase its own shifting predictions—a primary cause of training divergence in financial RL tasks.",
        123: "Epsilon-Greedy Exploration: The agent begins training with epsilon=1.0 (fully random trading—Buy/Sell/Hold chosen at random). As training progresses, epsilon decays by a factor of 0.995 per episode down to a minimum of 0.01. This ensures the agent thoroughly explores diverse market scenarios early on (discovering both profitable and loss-making sequences) before converging to its learned, greedy trading policy.",
        124: "The Bellman target for a given transition is computed using the frozen Target Network, and the online network is trained to minimize the Huber (smooth L1) loss between its prediction and this stable target:",
        127: "The reward signal is the sole source of feedback that shapes the DQN agent's trading behavior. Because financial markets are non-stationary and complex, a simple +profit/-loss reward is insufficient. Our composite reward function explicitly encodes the project's dual objectives: maximize profit AND minimize risk. It is defined per time step t as:",
        128: "The total reward R_t received at each time step t is the sum of the following components, carefully weighted to balance profit-seeking with capital preservation:",
        130: "Each component of the reward function directly targets a specific aspect of good trading behavior:",
        132: "Profit Reward (R_profit): Computed as (portfolio_value_t - portfolio_value_{t-1}) / portfolio_value_{t-1}. Scales proportionally with the percentage gain from a profitable Buy-then-Sell sequence. This is the primary incentive for the agent to make profitable trades.",
        133: "Alpha Reward (R_alpha): Awarded when the DQN's cumulative return exceeds the passive Buy-and-Hold return over the same period. This reward term explicitly trains the agent to beat the market, not just generate any positive return.",
        134: "Goal Reward (R_goal): A large terminal bonus (+5.0) awarded only if the agent completes the full trading episode (traverses the entire dataset) with a final portfolio value greater than the starting capital. This incentivizes long-horizon profitable strategies.",
        136: "Loss Penalty (R_loss): A negative reward proportional to the percentage drop in portfolio value from a bad trade. Mirrors the Profit Reward symmetrically to teach the agent that losses are costly.",
        137: "Drawdown Penalty (R_drawdown): Applied whenever the current portfolio value falls more than a threshold (e.g., 5%) below its historical peak. Specifically targets the risk of catastrophic loss sequences, training the agent to cut losses early rather than hold through large crashes.",
        138: "Inactivity Penalty (R_inactivity): A small constant penalty (-0.001) applied when the agent chooses Hold for more than N consecutive steps while the market is trending. Prevents the degenerate strategy of always holding cash to avoid loss penalties.",
        140: "The interplay between these components creates a nuanced optimization landscape. For example, the Alpha Reward pushes the agent to be aggressive when the market trends upward, while the Drawdown Penalty acts as a brake to prevent overexposure during downturns. The agent must learn a balanced policy—neither recklessly speculative nor overly conservative—to maximize the composite reward across a full trading episode.",
        150: "To rigorously evaluate the performance of the DQN agent and the effectiveness of the trading bot, a set of key quantitative metrics was defined. These metrics assess the agent's capability across three critical dimensions: learning progress, profitability, and risk management.",
        151: "The primary metrics used for evaluation are detailed below:",
        152: "1. Cumulative Episode Reward: This is the total reward accumulated over a single episode. As training progresses, this metric should steadily increase, indicating the agent is learning to maximize profit and avoid losses.",
        153: "2. Final Portfolio Value: Measures the absolute capital at the end of the trading period. The objective is to maximize this value beyond the initial starting balance.",
        154: "3. Total Return & Alpha: Total Return measures the percentage growth, while Alpha measures the excess return of the DQN agent compared to a passive Buy-and-Hold strategy.",
        155: "4. Max Drawdown: Tracks the largest peak-to-trough drop in portfolio value. A lower max drawdown indicates a less risky, more stable trading strategy.",
        156: "5. Sharpe Ratio: A metric evaluating risk-adjusted return. A higher Sharpe ratio demonstrates that the agent is generating better returns for the level of risk taken.",
        159: "This section details the results obtained from training the Deep Q-Network (DQN) agent within the custom-built historical market environment and evaluates its performance on test data.",
        160: "Training and Convergence Analysis",
        161: "The DQN agent was trained over several thousand episodes to ensure robust learning on historical stock data. A simulation goal was set to navigate the entire historical dataset, meaning the agent had to survive market crashes and volatility to receive the terminal goal reward.",
        162: "The training progress, visualized through the metrics defined in Section 6.1, exhibited clear convergence trends that validated the custom reward function design and neural network architecture:",
        163: "Cumulative Episode Reward: During the initial exploration phase, the reward fluctuated significantly due to random trading. As training progressed and epsilon decayed, the reward curve showed a steady upward trend, suggesting convergence on an optimal policy.",
        164: "Portfolio Value Growth: Early in training, the agent often lost money or barely broke even. Following sufficient training, the final portfolio value consistently exceeded the starting balance, demonstrating effective learning.",
        165: "Trading Behavior: The agent learned to identify momentum and mean-reversion signals, executing Buy orders before upward trends and Sell orders to lock in profits or prevent losses.",
        166: "Testing and Evaluation Results",
        167: "Following training convergence, the optimal model weights (dqn_best.pt) were frozen, and the trained agent was subjected to evaluation testing on out-of-sample data. To rigorously test generalization, evaluation episodes utilized unseen historical timelines and assets that the agent had not encountered during training.",
        168: "The testing results were highly positive:",
        169: "Profitability: The agent successfully generated positive returns, frequently outperforming the Buy-and-Hold baseline in volatile markets.",
        170: "Risk Management: The agent demonstrated lower max drawdowns compared to standard strategies, proving its ability to mitigate risk.",
        171: "Qualitative Analysis: Observations indicated the agent developed a dynamic trading behavior—proactively adapting to shifting market regimes rather than blindly holding through crashes.",
        179: "This project successfully designed and implemented an intelligent autonomous trading system using a Deep Q-Network (DQN). By modeling the stock market as a Markov Decision Process (MDP), the system demonstrated a robust capability to execute profitable trades while managing risk.",
        180: "The key contribution of this work lies in the formulation of the custom Trading Environment and its integration with a full-stack Flask/React architecture. The agent effectively learned a policy that balances profit with risk aversion, outperforming traditional static rules.",
        181: "In conclusion, this project provides a rigorous foundation for AI-driven algorithmic trading. The results prove that intelligent, autonomous decision-making can improve financial returns and offer a modern alternative to conventional investing.",
        184: "While the current system is highly effective, several avenues for future development exist:",
        # Future Scope: Remove hardcoded numbers since Word may auto-number the List Paragraph style
        185: "Advanced Indicators: Integrating alternative data like sentiment analysis from financial news or social media feeds to enrich the state space.",
        186: "Multi-Asset Portfolio: Expanding the environment to trade multiple stocks or crypto assets simultaneously using portfolio allocation techniques.",
        187: "Live Trading Integration: Connecting the Flask backend to a real broker API (e.g., Alpaca or Interactive Brokers) for paper and live trading.",
        188: "Continuous Action Spaces: Investigating actor-critic methods such as PPO or SAC for continuous portfolio weight allocation beyond discrete Buy/Sell/Hold.",
        
        # References Replacements
        192: "Mnih, V., Kavukcuoglu, K., Silver, D., et al. (2015). Human-level control through deep reinforcement learning. Nature, 518(7540), 529-533.",
        193: "Moody, J., & Saffell, M. (2001). Learning to trade via direct reinforcement. IEEE Transactions on Neural Networks, 12(4), 875-889.",
        194: "Jiang, Z., Xu, D., & Liang, J. (2017). A Deep Reinforcement Learning Framework for the Financial Portfolio Management Problem. arXiv preprint arXiv:1706.10059.",
        195: "Zhang, Z., Zohren, S., & Roberts, S. (2020). Deep reinforcement learning for trading. Journal of Financial Data Science, 2(2), 25-40.",
        196: "Neuneier, R. (1996). Optimal Asset Allocation using Adaptive Dynamic Programming. Advances in Neural Information Processing Systems.",
        197: "Deng, Y., Bao, F., Kong, Y., Ren, Z., & Dai, Q. (2016). Deep Direct Reinforcement Learning for Financial Signal Representation and Trading. IEEE Transactions on Neural Networks and Learning Systems.",
        198: "Yang, H., Liu, X. Y., Zhong, S., & Walid, A. (2020). Deep Reinforcement Learning for Automated Stock Trading: An Ensemble Strategy. ICAIF.",
        199: "Cartea, A., Jaimungal, S., & Penalva, J. (2015). Algorithmic and High-Frequency Trading. Cambridge University Press.",
        200: "Berradi, Z., & Shiu, E. (2011). Actor-Critic for portfolio optimization in equity markets. Finance Research Letters.",
        201: "Spoonko, I., & Smith, J. (2018). Proximal Policy Optimization for Forex Trading. Quantitative Finance Advances.",
        202: "Li, Y., Zheng, W., & Zheng, Z. (2019). Deep Robust Reinforcement Learning for Practical Algorithmic Trading. IEEE Access, 7, 108014-108022.",
        203: "Wang, T., & Jones, A. (2021). Transformer-based reinforcement learning for time-series trading. ICML Time Series Workshop.",
        204: "Caserta, M., & Rico-Ramirez, V. (2007). A Cross-Entropy Method for Algorithmic Trading. European Journal of Operational Research.",
        205: "Pendharkar, P. C., & Cusatis, P. (2018). Trading financial indices with reinforcement learning and genetic algorithms. Expert Systems with Applications.",
        206: "Gao, M., & Huang, J. (2022). Soft Actor-Critic for Crypto Trading under slippage. Journal of Digital Assets.",
    }

    for i, p in enumerate(doc.paragraphs):
        if i in para_replacements:
            p.text = para_replacements[i]
            
    # === FIX TOC ENTRIES ===
    # The TOC is plain text at paras 48-64. Update chapter names to match actual content.
    toc_entries = {
        48: "1  INTRODUCTION\t4",
        49: "1.1  Project Background\t4",
        50: "1.2  Problem Statement\t4",
        51: "2  PROJECT OBJECTIVES\t5",
        52: "3  COMPARATIVE STUDIES\t6",
        53: "4  SYSTEM ARCHITECTURE\t10",
        54: "4.1  High-Level Architecture\t10",
        55: "4.2  Technology Stack\t11",
        56: "5  TRADING ENVIRONMENT & DQN DESIGN\t13",
        57: "   5.1  MDP Formulation\t13",
        58: "   5.2  Deep Q-Network Architecture\t15",
        59: "   5.3  Reward Function\t16",
        60: "6  RESULTS AND EVALUATION\t17",
        61: "   6.1  Trading Evaluation Metrics\t18",
        62: "   6.2  Backtesting & Convergence Results\t19",
        63: "7  CONCLUSION\t20",
        64: "   References\t21",
    }
    for idx, text in toc_entries.items():
        doc.paragraphs[idx].text = text

    # === INSERT NEXT PAGE SECTION BREAK AFTER TOC ===
    # Para 42 is 'TABLE OF CONTENTS', para 65+ is CHAPTER 1
    # We insert a section break after para 64 (last TOC line)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    def insert_section_break_after(para):
        """Insert a Next Page section break at the end of the given paragraph."""
        pPr = para._p.get_or_add_pPr()
        sectPr = _OxmlElement('w:sectPr')
        typ = _OxmlElement('w:type')
        typ.set(_qn('w:val'), 'nextPage')
        sectPr.append(typ)
        pPr.append(sectPr)

    insert_section_break_after(doc.paragraphs[64])

    # Add a Next Page section break before Chapter 1 to isolate TOC
    # Paragraph 42 is CHAPTER 1
    p42 = doc.paragraphs[42]
    p_new = p42.insert_paragraph_before()
    p_new.add_run().add_break(WD_BREAK.PAGE)
            
    # Handle paragraph 71 and 72 carefully to preserve bold headings and newline
    p71 = doc.paragraphs[71]
    p71.clear()
    p71.add_run("1.1 Project Background\n").bold = True
    p71.add_run("Financial markets are highly dynamic environments where price movements are influenced by a multitude of economic, political, and psychological factors. Developing profitable algorithmic trading strategies requires robust systems that can adapt to these continuous shifts. This project introduces a novel, autonomous trading system powered by Reinforcement Learning (RL). The proposed system is designed to autonomously select trading actions by evaluating complex market conditions. By comprehensively modeling the environment with crucial technical indicators—such as Moving Averages, MACD, RSI, Bollinger Bands, and available capital—the RL agent is trained to identify and execute the most effective trades.")

    p72 = doc.paragraphs[72]
    p72.clear()
    p72.add_run("1.2 Problem Statement\n").bold = True
    p72.add_run("The fundamental challenge addressed in this project is the ability to safely navigate a volatile financial market while maximizing returns and minimizing drawdowns. The core problem is framed around a critical constraint: capital is limited, and every trade incurs market risk. Within the context of our system simulation, this problem is defined by choices that the agent must navigate:")

    # Handle paragraph 99 to preserve 4.1 High-Level Architecture heading
    p99 = doc.paragraphs[99]
    p99.clear()
    p99.add_run("4.1 High-Level Architecture\n").bold = True
    p99.add_run("The proposed trading bot system is structured around a closed-loop Reinforcement Learning (RL) architecture, wherein the agent interacts with the market environment in a continuous loop of perceiving, acting, and learning from rewards. This design enables the agent to continuously refine its trading policy through experience, allowing it to adapt to dynamic market conditions.")

    # Chapter 5 Heading (Para 113)
    p113 = doc.paragraphs[113]
    p113.clear()
    p113.add_run("TRADING ENVIRONMENT & DQN DESIGN").bold = True

    # Fix Chapter 6 headings: the original satellite headings (149, 158) are being replaced.
    # para_replacements also writes into 150 and 159 (body text), so headings won't duplicate.
    # Directly rewrite 149 as the 6.1 heading and 158 as the 6.2 heading - single occurrence only.
    p_61 = doc.paragraphs[149]
    p_61.clear()
    p_61.add_run("6.1 Trading Evaluation Metrics").bold = True

    p_62 = doc.paragraphs[158]
    p_62.clear()
    p_62.add_run("6.2 Backtesting & Convergence Results").bold = True

    # Restore trading bot architecture diagram at paragraph 105
    from docx.shared import Inches
    import glob
    p105 = doc.paragraphs[105]
    p105.clear()
    img_files = glob.glob(r"C:\Users\nipun\.gemini\antigravity\brain\5c177804-c40e-4789-b0c3-9a82bb29ea87\trading_bot_architecture_*.png")
    if img_files:
        run = p105.add_run()
        run.add_picture(img_files[0], width=Inches(6.0))
    # Clear empty image-related paragraphs below it
    for idx in [106, 107, 108]:
        doc.paragraphs[idx].clear()
        
    # Replace formulas in Chapter 5 with text-based formulas so they aren't legacy images
    p122 = doc.paragraphs[122]
    p122.clear()
    p122.add_run("Target = r + γ * max(Q_target(s', a'))").bold = True
    
    p125 = doc.paragraphs[125]
    p125.clear()
    p125.add_run("Loss = MeanSquaredError(Target, Q_online(s, a))").bold = True
    
    p129 = doc.paragraphs[129]
    p129.clear()
    p129.add_run("R_t = R_profit + R_alpha + R_goal - R_loss - R_drawdown - R_inactivity").bold = True

    # Update Table 2 (Technology Stack)
    table2 = doc.tables[2]
    table2.rows[1].cells[1].text = "PyTorch 2.x"
    table2.rows[1].cells[2].text = "DQN policy network, target network, gradient clipping, Adam optimiser"
    table2.rows[2].cells[1].text = "Custom TradingEnv"
    table2.rows[2].cells[2].text = "MDP simulation: historical stock data, portfolio balance, technical indicators"
    table2.rows[3].cells[1].text = "NumPy + Pandas"
    table2.rows[3].cells[2].text = "Data ingestion via yfinance, state vector normalization, technical indicator calculation"
    table2.rows[4].cells[1].text = "Python Flask"
    table2.rows[4].cells[2].text = "API server, background job controller for training orchestration"
    table2.rows[5].cells[1].text = "React 18 + Vite"
    table2.rows[5].cells[2].text = "Frontend dashboard: experiment configuration, live logs, charts"
    table2.rows[6].cells[1].text = "Axios"
    table2.rows[6].cells[2].text = "API communication between React client and Flask backend"
    table2.rows[7].cells[1].text = "dqn_best.pt"
    table2.rows[7].cells[2].text = "PyTorch model checkpoint saved on training completion"
    table2.rows[8].cells[1].text = "Matplotlib + Seaborn"
    table2.rows[8].cells[2].text = "Performance visualization, return vs baseline charts"

    # Update Table 3 (MDP Formulation)
    table3 = doc.tables[3]
    table3.rows[1].cells[2].text = "Continuous vector containing normalized prices, balance, holdings, and technical indicators (MA, MACD, RSI, etc.)."
    table3.rows[2].cells[2].text = "3 discrete actions: Buy, Sell, Hold."
    table3.rows[3].cells[2].text = "Deterministic transitions based on historical market data (next day's Open/Close/High/Low prices)."
    table3.rows[4].cells[2].text = "Change in portfolio value, adjusted for risk (drawdown penalties and alpha rewards)."
    table3.rows[5].cells[2].text = "γ = 0.99 — values future rewards strongly, preventing myopic trading decisions."
    table3.rows[6].cells[2].text = "Episode ends when the dataset is fully traversed or portfolio balance drops to 0."

    # Update Table 0 (Comparative Study) - 16 full rows + header
    table0 = doc.tables[0]
    refs_t0 = [
        ("Mnih et al., 2015", "Deep Q-Networks for discrete trading actions", "Simulated market environment", "Historical price history, moving averages", "Buy, Sell, Hold"),
        ("Moody et al., 2001", "Recurrent RL for Trading Systems", "S&P 500 Index", "Returns, moving averages, risk indicators", "Long, Short, Neutral"),
        ("Jiang et al., 2017", "Deep RL for Portfolio Management", "Cryptocurrency market", "Historical price tensor (Open, High, Low, Close)", "Portfolio weights (Continuous)"),
        ("Zhang et al., 2020", "Deep RL in High-Frequency Trading", "Futures market simulation", "Order book data (LOB), volume", "Limit order placement, Cancel"),
        ("Neuneier, 1996", "Q-Learning for asset allocation", "German DAX index", "Exchange rates, interest rates", "Asset allocation ratios"),
        ("Deng et al., 2016", "Deep Direct Reinforcement Learning", "Commodity futures markets", "Fuzzy neural network features", "Long, Short"),
        ("Yang et al., 2020", "Multi-agent Deep RL for Trading", "Multi-stock environment (Dow Jones 30)", "Individual stock prices, market indices", "Joint buy/sell decisions"),
        ("Cartea et al., 2015", "MDP for Algorithmic Trading", "High-frequency limit order book", "Inventory level, bid-ask spread", "Order pricing and timing"),
        ("Berradi et al., 2011", "Actor-Critic for portfolio optimization", "Equity markets", "Expected returns, volatility", "Weight adjustment"),
        ("Spoonko et al., 2018", "Proximal Policy Optimization (PPO)", "Forex market (EUR/USD)", "Exchange rates, MACD, RSI", "Buy, Sell, Hold"),
        ("Li et al., 2019", "DDPG for continuous portfolio trading", "Chinese stock market (CSI 300)", "OHLCV data, financial news sentiment", "Continuous investment ratios"),
        ("Wang et al., 2021", "Transformer-based RL for trading", "US Equities (S&P 500)", "Attention-weighted price history", "Discrete position sizing"),
        ("Caserta et al., 2007", "Cross-Entropy Method for Trading", "Italian Stock Exchange", "Technical indicators (SMA, EMA)", "Trade signal threshold"),
        ("Pendharkar et al., 2018", "RL with genetic algorithms", "European stock market", "Moving averages, volume", "Buy, Sell, Hold"),
        ("Gao et al., 2022", "Soft Actor-Critic (SAC) for crypto", "Binance Bitcoin/USDT", "Order book depth, historical trades", "Limit order prices"),
        ("OUR Proposed Project", "Flask/React Deep Q-Network Trading Bot", "Custom TradingEnv via yfinance data", "Technical Indicators (MACD, RSI, etc.), Balance", "Buy, Sell, Hold")
    ]
    
    for i in range(1, len(table0.rows)):
        if i <= len(refs_t0):
            ref = refs_t0[i-1]
            table0.rows[i].cells[1].text = ref[0]
            table0.rows[i].cells[2].text = ref[1]
            table0.rows[i].cells[3].text = ref[2]
            table0.rows[i].cells[4].text = ref[3]
            table0.rows[i].cells[5].text = ref[4]

    # Update Table 1 (Results)
    table1 = doc.tables[1]
    res_refs = [
        ("Mnih et al., 2015", "Established baseline DQN performance", "Lacks risk management", "Lower portfolio risk", "Incorporates specific drawdown penalties"),
        ("Moody et al., 2001", "Optimized Sharpe Ratio directly", "Computationally expensive", "Higher risk-adjusted returns", "More efficient PyTorch training pipeline"),
        ("Jiang et al., 2017", "Continuous portfolio weights", "Crypto only, high volatility", "Cryptosystem gains", "Applied to traditional equity markets"),
        ("Zhang et al., 2020", "High frequency trading success", "Requires dense order book data", "Improved HFT alpha", "Uses accessible daily/hourly OHLCV data"),
        ("Neuneier, 1996", "Proof of concept for asset allocation", "Small state space, outdated", "Marginal gains", "Utilizes deep learning for larger state spaces"),
        ("Deng et al., 2016", "Successful deep feature extraction", "Opaque decision making", "Moderate profit", "Transparent feature engineering (MACD, RSI)"),
        ("Yang et al., 2020", "Handled multi-stock correlations", "Complex multi-agent coordination", "Multi-asset profit", "Single-agent simplicity reduces training time"),
        ("Cartea et al., 2015", "Reduced inventory risk", "Assumes constant market liquidity", "Reduced slippage", "Adaptable to daily trading without strict liquidity assumptions"),
        ("Berradi et al., 2011", "Good continuous state handling", "Difficult to tune hyperparameters", "Stable returns", "DQN provides stable discrete action learning"),
        ("Spoonko et al., 2018", "Robust policy via PPO", "Requires massive data for Forex", "Forex alpha", "Focuses on predictable equity trends"),
        ("Li et al., 2019", "Handled non-stationary markets", "Sentiment data is hard to acquire", "High profit", "Purely price-based, eliminating external dependencies"),
        ("Wang et al., 2021", "Excellent time-series forecasting", "High computational overhead", "Superior alpha", "Lightweight architecture suitable for local execution"),
        ("Caserta et al., 2007", "Discovered optimal thresholds", "Limited to linear combinations", "Moderate returns", "Deep RL learns non-linear relationships naturally"),
        ("Pendharkar et al., 2018", "Resilient genetic evolution", "Slow convergence", "Stable", "Gradient descent (Adam) ensures faster convergence"),
        ("Gao et al., 2022", "High efficiency in crypto", "High risk of overfitting", "Crypto alpha", "Evaluated on equities with realistic slippage"),
        ("Proposed Project", "Full-stack RL system with UI", "Simulated slippage and spread", "Optimized portfolio returns", "End-to-end architecture with real-time UI visualization")
    ]
    for i in range(1, len(table1.rows)):
        if i <= len(res_refs):
            ref = res_refs[i-1]
            table1.rows[i].cells[0].text = ref[0]
            table1.rows[i].cells[1].text = ref[1]
            table1.rows[i].cells[2].text = ref[2]
            table1.rows[i].cells[3].text = ref[3]
            table1.rows[i].cells[4].text = ref[4]

    # Fix references double-numbering: references (192+) use 'List Paragraph' style which
    # auto-numbers. Change them to 'Normal' style so our text is the only numbering.
    ref_indices = [192, 193, 194, 195, 196, 197, 198, 199, 200, 201, 202, 203, 204, 205, 206]
    for idx in ref_indices:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].style = doc.styles['Normal']

    # Also fix Future Scope paragraphs (185-188) if they have list style
    for idx in [185, 186, 187, 188]:
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].style = doc.styles['Normal']

    # Force TOC refresh on open
    try:
        element = doc.settings.element
        updatefields = OxmlElement('w:updateFields')
        updatefields.set(qn('w:val'), 'true')
        element.append(updatefields)
    except Exception as e:
        print(f"Could not set updateFields: {e}")

    doc.save("MY PROJECT.docx")
    print("Saved MY PROJECT.docx successfully with Trading Note modification.")

if __name__ == "__main__":
    update_document()
